from docx import Document
from docx.enum.style import WD_STYLE_TYPE  # Import WD_STYLE_TYPE
import spacy
import re

class TextPreprocessor:
    def __init__(self):
        self.nlp = spacy.load("en_core_web_sm", disable=["tagger", "parser", "ner"])

    def preprocess(self, doc_path: str) -> str:
        doc = Document(doc_path)
        # self.remove_headings(doc)
        self.remove_content_lists(doc)
        self.remove_text_lists(doc)
        self.remove_headers_footers(doc)
        self.remove_red_text(doc)
        cleaned_text = self.extract_text(doc)
        cleaned_text = self.remove_extra_whitespace(cleaned_text) # Remove extra whitespace from the entire text
        cleaned_text = self.remove_numberings(cleaned_text)
        return cleaned_text

    # def remove_headings(self, doc: Document):
    #     for para in doc.paragraphs:
    #         if self.is_heading(para):
    #             para.clear()

    def is_heading(self, para) -> bool:
        doc = self.nlp(para.text.strip())
        return len(doc) <= 5 or bool(re.match(r"^\d+(\.\d+)*[A-Z]?(\s|$)", para.text.strip()))

    def remove_content_lists(self, doc: Document):
        """Remove content lists based on heading style and name."""
        paragraphs = doc.paragraphs
        remove_content = False  # Flag to indicate if we're in the content section

        for i, para in enumerate(paragraphs):
            if para.style.type == WD_STYLE_TYPE.PARAGRAPH: # Check if it is paragraph style, otherwise, it might be table or other element
                if "contents" in para.text.strip().lower() and para.style.name.lower().startswith("heading"): #Check if it is heading and contains "contents"
                    remove_content = True
                    para.clear()  # Clear the "Contents" heading itself
                elif remove_content and para.style.name.lower().startswith("heading"): # Check if it is the next heading
                    remove_content = False  # Stop removing
                elif remove_content:
                    para.clear()  # Clear the paragraph if it's within the content section

    def remove_text_lists(self, doc: Document):
        ref_pattern = re.compile(r"^\[\d+[A-Z]?\]\s*.+$", re.MULTILINE)
        for para in doc.paragraphs:
            para.text = ref_pattern.sub("", para.text)  # Remove references in place

    def remove_headers_footers(self, doc: Document):
        for section in doc.sections:
            for header in [section.header, section.footer]:
                if header:
                    for paragraph in header.paragraphs:
                        paragraph.clear()

    def remove_red_text(self, doc: Document):
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.color.rgb is not None:  # Check if color is set
                    if run.font.color.rgb == (255, 0, 0): # Check for red color (RGB)
                        run.clear() # Clear the red text run

    def extract_text(self, doc: Document) -> str:
        cleaned_text = ""
        for para in doc.paragraphs:
            cleaned_text += para.text + "\n"
        return cleaned_text

    def remove_extra_whitespace(self, text: str) -> str:
        text = re.sub(r"\n\s*\n", "\n", text) # Multiple newlines with single newline
        text = re.sub(r"\s+", " ", text) # Replace multiple spaces with single space
        return text.strip()
    
    def remove_numberings(self, text: str) -> str:
        # Remove patterns like "a)", "b)", ..., "1)", "2)" at the start of a line or after newlines
        cleaned_text = re.sub(r"(?m)^\s*[a-zA-Z0-9]+\)\s*", "", text)
        return cleaned_text


