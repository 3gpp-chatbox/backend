# Text Extraction from DOCX Files
from docx import Document
import os

def extract_text_from_docx(spec_name="TS_23_501"):
    """Extract text from 3GPP specification document"""
    # Define paths
    base_dir = "../3GPP_Documents"
    spec_dir = os.path.join(base_dir, spec_name)
    if not os.path.exists(spec_dir):
        print(f"Directory {spec_dir} does not exist")
        return ""
    
    # Look for the .docx file in the unzipped directory
    docx_files = [f for f in os.listdir(spec_dir) if f.endswith('.docx')]
    
    if not docx_files:
        print(f"No .docx file found in {spec_dir}")
        return ""
        
    # Use the first .docx file found
    docx_path = os.path.join(spec_dir, docx_files[0])
    print(f"Processing document: {docx_path}")
    
    try:
        # Load and process the document
        doc = Document(docx_path)
        
        # Extract text from all paragraphs
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
                
        print(f"Successfully extracted {len(text)} characters")
        return text
        
    except Exception as e:
        print(f"Error processing document: {str(e)}")
        return ""

if __name__ == "__main__":
    text = extract_text_from_docx()
    print("\nFirst 500 characters of extracted text:")
    print(text[:500])
