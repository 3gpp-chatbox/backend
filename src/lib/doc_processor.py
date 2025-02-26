"""
Document processor module for handling .docx specification files.
Provides functionality to read and extract text content from 3GPP specification documents.
"""

from pathlib import Path
import re
from typing import Dict, List, Optional
from docx import Document
from dataclasses import dataclass

@dataclass
class Section:
    """Represents a document section with its heading and content"""

    level: int
    heading: str
    content: List[Dict[str, str]]
    subsections: List["Section"]
    parent: Optional["Section"] = None

def load_document(file_path: str) -> Document:
    """
    Load the document from the given file path.

    Args:
        file_path (str): Path to the .docx file

    Returns:
        Document: The loaded document object
    """
    try:
        print(f"Loading document from {file_path}")
        doc = Document(file_path)
        print("Document loaded successfully")
        return doc

    except FileNotFoundError:
        print(f"File not found: {file_path}")
        return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def extract_paragraphs(doc: Document) -> List[Dict[str, any]]:
    """
    Extract paragraphs from the document with their styles.

    Args:
        doc (Document): The loaded document object

    Returns:
        List[Dict]: List of paragraphs, each containing:
            - text: The paragraph text
            - style: The paragraph style name
            - level: The heading level (if applicable)
    """
    paragraphs = []
    for para in doc.paragraphs:
        if not para.text.strip():  # Skip empty paragraphs
            continue

        # Get paragraph style information
        style_name = para.style.name if para.style else "Normal"
        level = None

        # Check if it's a heading and get its level
        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                pass

        paragraphs.append(
            {
                "text": text_cleaner(para.text),
                "style": text_cleaner(style_name),
                "level": level,
            }
        )

    return paragraphs

def extract_section_tree(doc: Document, max_heading_level: int = 8) -> List[Section]:
    """
    Extract document content as a tree of sections based on heading hierarchy.
    This allows for processing the document in meaningful chunks based on its structure.

    Args:
        doc (Document): The loaded document object
        max_heading_level (int): Maximum heading level to consider for sectioning

    Returns:
        List[Section]: List of top-level sections, each containing their subsections
    """
    print("Extracting section tree")
    if doc is None:
        print("Document is None!")
        return []

    # Initialize variables
    current_sections = [None] * (max_heading_level + 1)  # Track current section at each level
    current_content = []  # Accumulate content for the current deepest section
    top_level_sections = []  # Store all level 1 sections
    max_chunk_size = 2000

    # Extract paragraphs from the document
    paragraphs = extract_paragraphs(doc)

    for para in paragraphs:
        level = para.get("level")

        # Check if the paragraph is a heading (level is not None and within max_heading_level)
        # Also ensure the heading text starts with a number (e.g., "4.1"), which may include decimals
        if level is not None and level <= max_heading_level and para["text"][0].isdigit():
            # Create a new section for the heading
            new_section = Section(
                level=level,
                heading=para["text"].strip().replace(" ", "_"),  # Replace spaces with underscores for consistency
                content=[],
                subsections=[],
                parent=current_sections[level - 1] if level > 1 else None,  # Link to parent section if not level 1
            )

            # Find the deepest non-None section (highest level) to assign current_content
            deepest_section = None
            for i in range(max_heading_level, 0, -1):
                if current_sections[i] is not None:
                    deepest_section = current_sections[i]
                    break
            if deepest_section is not None:
                deepest_section.content.extend(current_content)  # Add content to the deepest section

            # Update the section hierarchy
            if level > 1 and current_sections[level - 1] is not None:
                current_sections[level - 1].subsections.append(new_section)  # Link new section to parent

            current_sections[level] = new_section  # Set the new section at its level
            current_content = []  # Reset content after assigning it

            # Clear all lower-level sections (reset hierarchy below the current level)
            for i in range(level + 1, max_heading_level + 1):
                current_sections[i] = None

            # If this is a level 1 section, add it to top_level_sections
            if level == 1:
                top_level_sections.append(new_section)
        else:
            # Accumulate content for the current deepest section
            if (
                current_content
                and len(current_content[-1]) + len(para["text"]) < max_chunk_size
                and para.get("level") is None
            ):
                # Combine with previous text if conditions are met (within chunk size and not a heading)
                current_content[-1] = current_content[-1] + " " + para["text"]
            else:
                current_content.append(para["text"])  # Add new content block

    # Find the deepest current section to assign remaining content
    deepest_section = None
    for i in range(max_heading_level, 0, -1):
        if current_sections[i] is not None:
            deepest_section = current_sections[i]
            break
    if deepest_section is not None:
        deepest_section.content.extend(current_content)  # Add remaining content to the deepest section

    return top_level_sections

def text_cleaner(text: str) -> str:
    """
    Clean and normalize text while preserving specific patterns and cases.

    The function performs the following operations:
    1. Replaces non-breaking spaces and tabs with regular spaces
    2. Preserves version numbers (e.g., 1.2, 5.5.1.2)
    3. Removes punctuation marks (.,!?:) when they appear at the end of words
    4. Normalizes whitespace
    5. Performs case normalization with specific preservation rules:
        - Preserves words containing numbers (e.g., "5G")
        - Preserves words with multiple uppercase letters (e.g., "IP", "NSSAI")
        - Preserves words containing special characters (e.g., "S-NSSAI(s)")
        - Converts all other words to lowercase

    Args:
        text (str): Input text to clean

    Returns:
        str: Cleaned and normalized text with preserved patterns
    """
    if not text:
        return ""

    # Replace non-breaking spaces and tabs with regular spaces
    text = text.replace("\xa0", " ").replace("\t", " ")

    # Preserve version numbers (e.g., 1.2, 5.5.1.2)
    # This pattern matches two or more numbers separated by dots
    text = re.sub(r"\b\d+\.\d+(?:\.\d+)*\b", lambda x: f"__{x.group(0)}__", text)

    # Remove punctuation marks only when they appear at the end of words
    text = re.sub(r"([.,!?:])\s", " ", text)
    text = re.sub(r"([.,!?:])$", "", text)

    text = text.replace("__", "")

    # Normalize whitespace
    text = " ".join(text.split())

    # Process each word for case normalization
    words = text.split()
    normalized_words = []
    for word in words:
        # Preserve if:
        # 1. Contains numbers
        # 2. Contains more than one uppercase letter
        # 3. Contains special characters (excluding common punctuation)
        if (
            any(c.isdigit() for c in word)  # Has numbers
            or sum(1 for c in word if c.isupper()) > 1  # Multiple uppercase
            or any(c for c in word if not c.isalnum())
        ):  # Has symbols
            normalized_words.append(word)
        else:
            normalized_words.append(word.lower())

    return " ".join(normalized_words)

def remove_sections(file_path: str, excluded_sections: List[str]) -> str:
    """
    Remove specified sections from a document and save the stripped version.

    Args:
        file_path (str): Path to the source document
        excluded_sections (List[str]): List of section names to remove from the document

    Returns:
        str: Path to the saved stripped document in the data/stripped directory

    This function removes all paragraphs that are part of the excluded sections and
    saves the resulting document with the same name in the data/stripped directory.
    """
    doc = load_document(file_path)

    remove = False
    for para in doc.paragraphs:
        para.text = text_cleaner(para.text)
        if para.style.name.startswith("Heading") and (
            any(word.lower() in para.text.lower() for word in excluded_sections)
            or not para.text.strip()[0].isdigit()
        ):
            print(f"Removing section: {para.text}")
            remove = True
        elif para.style.name.startswith("Heading") and remove:
            remove = False
        if remove:
            para.clear()
            p = para._element
            p.getparent().remove(p)

    # Extract original filename and save stripped version
    doc_name = Path(file_path).name
    save_path = f"data/stripped/{doc_name}"
    doc.save(save_path)
    return save_path


def extract_table_of_contents(section_tree: List[Section]) -> str:
    """
    Generate a hierarchical table of contents string from a section tree.
    
    Args:
        section_tree (List[Section]): List of top-level sections
        
    Returns:
        str: A formatted string containing the table of contents with proper indentation
        showing the hierarchical structure of the document.
        
    Example output:
        1. Introduction
            1.1 Overview
            1.2 Scope
                1.2.1 Applications
        2. Requirements
            2.1 System Requirements
            2.2 User Requirements
    """
    def _process_section(section: Section, toc_lines: List[str]):
        """Helper function to recursively process sections and their subsections"""
        # Calculate indentation based on section level (4 spaces per level)
        indent = "    " * (section.level - 1)
        
        # Add the current section heading with proper indentation
        # The heading already contains the section number since we preserved it in extract_section_tree
        toc_lines.append(f"{indent}{section.heading}")
        
        # Process all subsections recursively
        for subsection in section.subsections:
            _process_section(subsection, toc_lines)
    
    toc_lines = []
    
    # Process each top-level section
    for section in section_tree:
        _process_section(section, toc_lines)
    
    # Join all lines with newlines to create the final TOC string
    return "\n".join(toc_lines)
