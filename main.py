from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from docx import Document
import os
import sqlite3
import json
from typing import Dict, List
from extract_procedures.graph_extractor import GraphExtractor

# Create FastAPI app instance
app = FastAPI(
    title="3GPP Procedures API",
    description="API for managing and visualizing 3GPP procedures",
    version="1.0.0"
)

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # React frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize GraphExtractor
extractor = GraphExtractor(db_path="nas_chunks.db")

@app.get("/")
async def root():
    return {"message": "Welcome to 3GPP Procedures API"}

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

@app.get("/process-directory/{dir_name}")
async def process_directory(dir_name: str):
    """
    Process all .docx files in a specified directory
    """
    try:
        base_dir = "3GPP_Documents"
        spec_dir = os.path.join(base_dir, dir_name)
        
        if not os.path.exists(spec_dir):
            raise HTTPException(status_code=404, detail=f"Directory {spec_dir} does not exist")
        
        docx_files = [f for f in os.listdir(spec_dir) if f.endswith('.docx')]
        
        if not docx_files:
            raise HTTPException(status_code=404, detail=f"No .docx files found in {spec_dir}")
        
        results = []
        for docx_file in docx_files:
            docx_path = os.path.join(spec_dir, docx_file)
            print(f"Processing document: {docx_path}")
            
            try:
                doc = Document(docx_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                nodes, edges = extractor.process_text_in_chunks(text)
                
                # Get summary for this file
                summary = {
                    "filename": docx_file,
                    "text_length": len(text),
                    "nodes_count": {k: len(v) for k, v in nodes.items()},
                    "edges_count": len(edges),
                    "nodes": {k: list(v) for k, v in nodes.items()},
                    "edges": edges
                }
                results.append(summary)
                
            except Exception as e:
                print(f"Error processing {docx_file}: {str(e)}")
                results.append({
                    "filename": docx_file,
                    "error": str(e)
                })
        
        return JSONResponse({
            "directory": dir_name,
            "total_files": len(docx_files),
            "processed_files": results
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/summary")
async def get_summary():
    """
    Get a summary of all processed data
    """
    try:
        # Get database summary
        conn = sqlite3.connect(extractor.db_path)
        c = conn.cursor()
        c.execute('SELECT COUNT(*) FROM chunk_results')
        total_chunks = c.fetchone()[0]
        conn.close()
        
        # Get current nodes and edges
        nodes = {k: list(v) for k, v in extractor.nas_nodes.items()}
        edges = extractor.nas_edges
        
        # Create detailed summary
        summary = {
            "database_stats": {
                "total_chunks": total_chunks
            },
            "extraction_stats": {
                "total_nodes": sum(len(v) for v in nodes.values()),
                "nodes_by_type": {k: len(v) for k, v in nodes.items()},
                "total_edges": len(edges)
            },
            "nodes": nodes,
            "edges": edges
        }
        
        # Print summary to console
        print("\nNAS Graph Elements Summary:")
        for node_type, node_list in nodes.items():
            print(f"\n{node_type.capitalize()} ({len(node_list)}):")
            print('\n'.join(f"- {node}" for node in node_list))
        
        print("\nEdges:")
        for edge in edges:
            print(f"{edge['source']} --[{edge['relationship']}]--> {edge['target']} (Procedure: {edge['procedure']})")
        
        return JSONResponse(summary)
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/view-database")
async def view_database():
    """
    View all stored chunks and their extracted information from the database
    """
    try:
        conn = sqlite3.connect(extractor.db_path)
        c = conn.cursor()
        c.execute('SELECT * FROM chunk_results')
        rows = c.fetchall()
        conn.close()
        
        results = []
        for row in rows:
            results.append({
                "id": row[0],
                "chunk_text": row[1],
                "nodes": json.loads(row[2]),
                "edges": json.loads(row[3])
            })
        
        return JSONResponse({
            "total_chunks": len(results),
            "chunks": results
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/extract-procedures")
async def extract_procedures(file: UploadFile = File(...)):
    """
    Extract procedures from an uploaded .docx file
    """
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    try:
        # Save the uploaded file temporarily
        temp_path = f"temp_{file.filename}"
        print(f"Processing file: {file.filename}")
        
        with open(temp_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        # Extract text from the document
        doc = Document(temp_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        print(f"Extracted text length: {len(text)} characters")
        
        # Process the text using GraphExtractor
        print("Starting text processing...")
        nodes, edges = extractor.process_text_in_chunks(text)
        
        # Print summary of extracted data
        print("\nExtraction Summary:")
        print(f"Total nodes: {sum(len(v) for v in nodes.values())}")
        for node_type, values in nodes.items():
            print(f"- {node_type}: {len(values)} items")
        print(f"Total edges: {len(edges)}")
        
        # Convert sets to lists for JSON serialization
        nodes_json = {k: list(v) for k, v in nodes.items()}
        
        # Clean up temporary file
        os.remove(temp_path)
        print("Processing completed successfully")
        
        return JSONResponse({
            "nodes": nodes_json,
            "edges": edges
        })
        
    except Exception as e:
        print(f"Error during processing: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/procedures")
async def get_procedures():
    """
    Get all extracted procedures from the database
    """
    try:
        return {
            "nodes": {k: list(v) for k, v in extractor.nas_nodes.items()},
            "edges": extractor.nas_edges
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/procedure/{procedure_id}")
async def get_procedure_by_id(procedure_id: int):
    """
    Get a specific procedure by its ID from the database
    """
    try:
        conn = sqlite3.connect(extractor.db_path)
        c = conn.cursor()
        c.execute('SELECT chunk_text, nodes, edges FROM chunk_results WHERE id = ?', (procedure_id,))
        result = c.fetchone()
        conn.close()
        
        if result is None:
            raise HTTPException(status_code=404, detail=f"Procedure with id {procedure_id} not found")
            
        chunk_text, nodes_json, edges_json = result
        
        return JSONResponse({
            "id": procedure_id,
            "chunk_text": chunk_text,
            "nodes": json.loads(nodes_json),
            "edges": json.loads(edges_json)
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/procedures/types")
async def get_procedure_types():
    """
    Get all available procedure types from the database
    """
    try:
        conn = sqlite3.connect(extractor.db_path)
        c = conn.cursor()
        
        # First, get all unique procedures from the database
        c.execute('''
            SELECT DISTINCT json_extract(nodes, '$.procedures[0]') as proc
            FROM chunk_results 
            WHERE json_extract(nodes, '$.procedures[0]') IS NOT NULL
        ''')
        procedures = [row[0] for row in c.fetchall() if row[0]]
        
        # If no procedures found in database, return predefined types
        if not procedures:
            predefined_types = [
                'NAS Registration Procedure',
                'NAS Authentication Procedure',
                'LTE Attach Procedure',
                'LTE Detach Procedure',
                '5G Registration Procedure',
                '5G PDU Session Procedure'
            ]
            return JSONResponse({
                "types": predefined_types
            })
        
        # Categorize procedures based on keywords
        categorized = {
            'NAS': [],
            'LTE': [],
            '5G': []
        }
        
        for proc in procedures:
            proc_lower = proc.lower()
            if any(kw in proc_lower for kw in ['nas', '5gmm', 'emm']):
                categorized['NAS'].append(proc)
            elif any(kw in proc_lower for kw in ['eps', 'lte', 'emm', 'esm']):
                categorized['LTE'].append(proc)
            elif any(kw in proc_lower for kw in ['5g', '5gs', '5gmm', '5gsm']):
                categorized['5G'].append(proc)
        
        # Format the results
        result = []
        for category, procs in categorized.items():
            for proc in procs:
                result.append(f"{category} {proc}")
        
        return JSONResponse({
            "types": sorted(result) if result else predefined_types
        })
        
    except Exception as e:
        print(f"Error getting procedure types: {str(e)}")
        # Return predefined types on error
        predefined_types = [
            'NAS Registration Procedure',
            'NAS Authentication Procedure',
            'LTE Attach Procedure',
            'LTE Detach Procedure',
            '5G Registration Procedure',
            '5G PDU Session Procedure'
        ]
        return JSONResponse({
            "types": predefined_types
        })
    finally:
        conn.close()

@app.get("/procedures/by-type/{procedure_type}")
async def get_procedures_by_type(procedure_type: str):
    """
    Get all procedures of a specific type
    """
    try:
        conn = sqlite3.connect(extractor.db_path)
        c = conn.cursor()
        c.execute('SELECT id, chunk_text, nodes, edges FROM chunk_results')
        rows = c.fetchall()
        conn.close()
        
        matching_procedures = []
        for row in rows:
            id, chunk_text, nodes_json, edges_json = row
            nodes = json.loads(nodes_json)
            if 'procedures' in nodes and any(
                proc.lower() == procedure_type.lower() 
                for proc in nodes['procedures']
            ):
                matching_procedures.append({
                    "id": id,
                    "chunk_text": chunk_text,
                    "nodes": nodes,
                    "edges": json.loads(edges_json)
                })
        
        return JSONResponse({
            "total": len(matching_procedures),
            "procedures": matching_procedures
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/procedures/flow/{procedure_type}")
async def get_procedure_flow(procedure_type: str):
    """
    Get the complete flow for a specific procedure type with sub-procedures
    """
    try:
        conn = sqlite3.connect(extractor.db_path)
        c = conn.cursor()
        c.execute('SELECT id, chunk_text, nodes, edges FROM chunk_results')
        rows = c.fetchall()
        conn.close()
        
        procedure_lower = procedure_type.lower()
        
        # Collect all nodes and edges related to this procedure
        all_nodes = {
            "states": set(),
            "messages": set(),
            "procedures": set(),
            "entities": set()
        }
        all_edges = []
        relevant_chunks = []
        
        for row in rows:
            id, chunk_text, nodes_json, edges_json = row
            nodes = json.loads(nodes_json)
            edges = json.loads(edges_json)
            
            # Check if this chunk is related to our procedure
            is_relevant = False
            if 'procedures' in nodes:
                for proc in nodes['procedures']:
                    proc_lower = proc.lower()
                    # Match both main procedure and potential sub-procedures
                    if (procedure_lower in proc_lower or
                        any(term in proc_lower for term in [
                            'initial', 'periodic', 'emergency', 'combined',
                            'ue-initiated', 'network-initiated', 'network-triggered'
                        ])):
                        is_relevant = True
                        break
            
            if is_relevant:
                # Add nodes
                for node_type in all_nodes:
                    if node_type in nodes:
                        all_nodes[node_type].update(nodes[node_type])
                # Add edges
                all_edges.extend(edges)
                # Store chunk info
                relevant_chunks.append({
                    "id": id,
                    "text": chunk_text
                })
        
        # If no data found in database, provide default flow data for predefined procedures
        if not all_edges and not any(len(nodes) > 0 for nodes in all_nodes.values()):
            # Default flows for predefined procedures
            DEFAULT_FLOWS = {
                'lte attach procedure': {
                    'nodes': {
                        'states': ['EMM-DEREGISTERED', 'EMM-REGISTERED', 'EMM-COMMON-PROCEDURE-INITIATED'],
                        'messages': ['Attach Request', 'Authentication Request', 'Authentication Response', 
                                   'Security Mode Command', 'Security Mode Complete', 'ESM Information Request',
                                   'ESM Information Response', 'Attach Accept', 'Attach Complete'],
                        'procedures': ['LTE Attach Procedure', 'Authentication Procedure', 'Security Mode Procedure'],
                        'entities': ['UE', 'MME', 'HSS', 'SGW']
                    },
                    'edges': [
                        {'source': 'EMM-DEREGISTERED', 'target': 'Attach Request', 'relationship': 'sends'},
                        {'source': 'Attach Request', 'target': 'EMM-COMMON-PROCEDURE-INITIATED', 'relationship': 'triggers'},
                        {'source': 'EMM-COMMON-PROCEDURE-INITIATED', 'target': 'Authentication Request', 'relationship': 'initiates'},
                        {'source': 'Authentication Request', 'target': 'Authentication Response', 'relationship': 'receives'},
                        {'source': 'Authentication Response', 'target': 'Security Mode Command', 'relationship': 'triggers'},
                        {'source': 'Security Mode Command', 'target': 'Security Mode Complete', 'relationship': 'receives'},
                        {'source': 'Security Mode Complete', 'target': 'ESM Information Request', 'relationship': 'triggers'},
                        {'source': 'ESM Information Request', 'target': 'ESM Information Response', 'relationship': 'receives'},
                        {'source': 'ESM Information Response', 'target': 'Attach Accept', 'relationship': 'triggers'},
                        {'source': 'Attach Accept', 'target': 'Attach Complete', 'relationship': 'receives'},
                        {'source': 'Attach Complete', 'target': 'EMM-REGISTERED', 'relationship': 'transitions'}
                    ]
                },
                'nas registration procedure': {
                    'nodes': {
                        'states': ['5GMM-DEREGISTERED', '5GMM-REGISTERED', '5GMM-COMMON-PROCEDURE-INITIATED'],
                        'messages': ['Registration Request', 'Authentication Request', 'Authentication Response',
                                   'Security Mode Command', 'Security Mode Complete', 'Registration Accept',
                                   'Registration Complete'],
                        'procedures': ['NAS Registration Procedure', 'Authentication Procedure', 'Security Mode Procedure'],
                        'entities': ['UE', 'AMF', 'AUSF', 'UDM']
                    },
                    'edges': [
                        {'source': '5GMM-DEREGISTERED', 'target': 'Registration Request', 'relationship': 'sends'},
                        {'source': 'Registration Request', 'target': '5GMM-COMMON-PROCEDURE-INITIATED', 'relationship': 'triggers'},
                        {'source': '5GMM-COMMON-PROCEDURE-INITIATED', 'target': 'Authentication Request', 'relationship': 'initiates'},
                        {'source': 'Authentication Request', 'target': 'Authentication Response', 'relationship': 'receives'},
                        {'source': 'Authentication Response', 'target': 'Security Mode Command', 'relationship': 'triggers'},
                        {'source': 'Security Mode Command', 'target': 'Security Mode Complete', 'relationship': 'receives'},
                        {'source': 'Security Mode Complete', 'target': 'Registration Accept', 'relationship': 'triggers'},
                        {'source': 'Registration Accept', 'target': 'Registration Complete', 'relationship': 'receives'},
                        {'source': 'Registration Complete', 'target': '5GMM-REGISTERED', 'relationship': 'transitions'}
                    ]
                }
            }
            
            # Check if we have a default flow for this procedure
            for key, flow in DEFAULT_FLOWS.items():
                if key in procedure_lower:
                    all_nodes = {k: set(v) for k, v in flow['nodes'].items()}
                    all_edges = flow['edges']
                    relevant_chunks = [{
                        "id": 0,
                        "text": f"Default flow diagram for {procedure_type}"
                    }]
                    break
        
        # Convert sets to lists for JSON
        nodes_json = {k: sorted(list(v)) for k, v in all_nodes.items()}
        
        return JSONResponse({
            "procedure_type": procedure_type,
            "nodes": nodes_json,
            "edges": all_edges,
            "relevant_chunks": relevant_chunks,
            "total_chunks": len(relevant_chunks)
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# Add new endpoint for sub-procedures
@app.get("/procedures/{procedure_type}/sub-procedures")
async def get_sub_procedures(procedure_type: str):
    """
    Get available sub-procedures for a specific procedure type
    """
    try:
        # Define sub-procedure mappings based on 3GPP specifications
        SUB_PROCEDURE_MAPPINGS = {
            'Registration': [
                'Initial Registration',
                'Periodic Registration',
                'Emergency Registration',
                'Mobility Registration'
            ],
            'Authentication': [
                'Primary Authentication',
                'EAP Authentication',
                'Secondary Authentication'
            ],
            'Security Mode Control': [
                'NAS Security Mode Control',
                'AS Security Mode Control'
            ],
            'PDU Session': [
                'PDU Session Establishment',
                'PDU Session Modification',
                'PDU Session Release'
            ],
            'Service Request': [
                'UE Triggered Service Request',
                'Network Triggered Service Request',
                'Emergency Service Request'
            ]
        }
        
        # Extract main procedure from procedure_type
        main_proc = procedure_type.split()[-1]
        
        if main_proc in SUB_PROCEDURE_MAPPINGS:
            return JSONResponse({
                "sub_procedures": SUB_PROCEDURE_MAPPINGS[main_proc]
            })
        else:
            return JSONResponse({
                "sub_procedures": []
            })
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
